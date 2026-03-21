"""Converts an AAR markdown report to a formatted .docx document.

Formatting is driven by a TemplateConfig parsed from a .docx template.
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips, Emu

from .template_config import TemplateConfig, FormatSpec


# Inline labels that get bold formatting with normal-weight content after
_INLINE_LABELS = (
    "PERSONNEL:", "EQUIPMENT:", "LOGISTICS:",
    "COMPOSITION:", "STRENGTH:", "CASUALTIES:",
    "Accuracy of Pre-Mission Intelligence:",
    "New Intelligence Gathered During The Operation:",
    "SUSTAINMENT:", "IMPROVEMENT:", "CONCLUSION:",
)

# Regex for section headers: "1. GENERAL INFORMATION ..."
_SECTION_RE = re.compile(r"^(\d+)\.\s+(.+)")

# Regex for memo fields: "TO:    value"
_MEMO_RE = re.compile(r"^(TO|FROM|SUBJECT|REF):\s*(.*)", re.IGNORECASE)

_ALIGN_MAP = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
}


class DocxConverter:
    """Converts an AAR markdown report to a formatted .docx
    using formatting from a TemplateConfig."""

    def __init__(
        self, report_text: str, template: TemplateConfig | None = None
    ) -> None:
        self._lines = report_text.strip().splitlines()
        self._doc = Document()
        self._t = template or TemplateConfig.default()
        self._font = self._t.default_font
        self._setup_defaults()

    def save(self, path: str | Path) -> None:
        """Parse the report and save to a .docx file."""
        self._add_header()
        self._build()
        self._add_footer()
        self._doc.save(str(path))

    # ------------------------------------------------------------------
    # Document setup
    # ------------------------------------------------------------------

    def _setup_defaults(self) -> None:
        """Set document-wide defaults from template config."""
        t = self._t
        style = self._doc.styles["Normal"]
        style.font.name = self._font
        style.font.size = Pt(t.default_size_pt)
        style.paragraph_format.space_before = Twips(t.default_space_before)
        style.paragraph_format.space_after = Twips(t.default_space_after)

        body_ls = t.body_format.line_spacing_twips
        if body_ls:
            style.paragraph_format.line_spacing = Twips(body_ls)

        margins = t.page_margins
        for section in self._doc.sections:
            section.top_margin = Twips(margins.get("top", 1440))
            section.bottom_margin = Twips(margins.get("bottom", 1440))
            section.left_margin = Twips(margins.get("left", 1440))
            section.right_margin = Twips(margins.get("right", 1440))
            section.header_distance = Twips(margins.get("header", 720))
            section.footer_distance = Twips(margins.get("footer", 720))

    # ------------------------------------------------------------------
    # Header
    # ------------------------------------------------------------------

    def _add_header(self) -> None:
        """Add the banner image to the document header."""
        if not self._t.banner_image:
            return
        section = self._doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        image_stream = io.BytesIO(self._t.banner_image)
        width = Emu(self._t.banner_width_emu) if self._t.banner_width_emu else None
        height = Emu(self._t.banner_height_emu) if self._t.banner_height_emu else None
        run.add_picture(image_stream, width=width, height=height)

    # ------------------------------------------------------------------
    # Footer
    # ------------------------------------------------------------------

    def _add_footer(self) -> None:
        """Add 'Page X of Y' footer."""
        section = self._doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = p.add_run("Page ")
        run.font.name = self._font
        run.font.size = Pt(10)
        run.italic = True

        self._add_field(p, "PAGE", italic=True)

        run = p.add_run(" of ")
        run.font.name = self._font
        run.font.size = Pt(10)
        run.italic = True

        self._add_field(p, "NUMPAGES", italic=True)

    def _add_field(self, paragraph, field_code: str, italic: bool = False) -> None:
        """Insert a Word field code (PAGE, NUMPAGES, etc.)."""
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run1 = paragraph.add_run()
        run1.font.name = self._font
        run1.font.size = Pt(10)
        run1.italic = italic
        run1._element.append(fld_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {field_code} "
        run2 = paragraph.add_run()
        run2.font.name = self._font
        run2.font.size = Pt(10)
        run2.italic = italic
        run2._element.append(instr)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run3 = paragraph.add_run()
        run3.font.name = self._font
        run3.font.size = Pt(10)
        run3.italic = italic
        run3._element.append(fld_end)

    # ------------------------------------------------------------------
    # Main builder
    # ------------------------------------------------------------------

    def _build(self) -> None:
        """Parse the report lines and build the document."""
        t = self._t
        i = 0
        n = len(self._lines)
        in_conclusion = False
        in_header = True
        in_memo = False
        last_section_num = len(t.sections)

        while i < n:
            line = self._lines[i]
            stripped = line.strip()

            if not stripped:
                if not in_header or in_memo:
                    self._doc.add_paragraph()
                i += 1
                continue

            # --- Title block ---
            if stripped == t.unit_name and i < 3:
                self._add_formatted_para(stripped, t.title_format)
                i += 1
                continue

            if stripped == t.unit_subline:
                self._add_formatted_para(stripped, t.subtitle_format)
                i += 1
                continue

            if ("AFTER ACTION REPORT" in stripped.upper()
                    and not _MEMO_RE.match(stripped)):
                self._add_formatted_para(stripped, t.op_line_format)
                i += 1
                continue

            # Fallback: "OPERATION" line merged with section header
            if stripped.startswith("OPERATION "):
                section_in_line = re.search(r"(\d+)\.\s+([A-Z].*)", stripped)
                if section_in_line:
                    op_part = stripped[:section_in_line.start()].strip()
                    sec_part = stripped[section_in_line.start():]
                    if op_part:
                        self._add_formatted_para(
                            op_part + " - AFTER ACTION REPORT", t.op_line_format
                        )
                    self._add_formatted_para(sec_part, t.section_header_format)
                    in_header = False
                else:
                    self._add_formatted_para(stripped, t.op_line_format)
                i += 1
                continue

            # Date line
            if re.match(r"^\[?\d{4}[/-]\d{2}[/-]\d{2}\]?$", stripped):
                p = self._doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Twips(480)
                body_ls = t.body_format.line_spacing_twips
                if body_ls:
                    p.paragraph_format.line_spacing = Twips(body_ls)
                run = p.add_run(stripped)
                run.font.name = self._font
                run.font.size = Pt(t.default_size_pt)
                i += 1
                continue

            # Location line (short, uppercase, before memo)
            if (in_header and not in_memo
                    and len(stripped) < 60
                    and not any(c.islower() for c in stripped.replace(" ", ""))
                    and not stripped[0].isdigit()
                    and stripped != t.unit_name
                    and not stripped.startswith("OPERATION")
                    and not stripped.startswith("MEMORANDUM")):
                self._add_formatted_para(stripped, t.op_line_format)
                i += 1
                continue

            # MEMORANDUM FOR
            if stripped == "MEMORANDUM FOR":
                self._add_formatted_para(stripped, t.memo_format)
                self._doc.add_paragraph()
                in_memo = True
                i += 1
                continue

            # Memo fields
            memo_match = _MEMO_RE.match(stripped)
            if memo_match and in_memo:
                label = memo_match.group(1).upper() + ":"
                value = memo_match.group(2).strip()
                self._add_memo_line(label, value)
                if label == "REF:":
                    in_memo = False
                    self._doc.add_paragraph()
                i += 1
                continue

            in_header = False

            # Section headers
            section_match = _SECTION_RE.match(stripped)
            if section_match:
                section_num = int(section_match.group(1))
                in_conclusion = section_num >= last_section_num
                self._add_formatted_para(stripped, t.section_header_format)
                i += 1
                continue

            # Bullet items
            if stripped.startswith("* "):
                self._add_bullet_item(stripped[2:])
                i += 1
                continue

            # Inline labels
            label_found = False
            for label in _INLINE_LABELS:
                if stripped.startswith(label):
                    content = stripped[len(label):].strip()
                    self._add_labeled_line(label, content)
                    label_found = True
                    break
            if label_found:
                i += 1
                continue

            # Signature block
            if in_conclusion and self._is_signature_line(stripped):
                sig_lines = []
                while i < n:
                    sl = self._lines[i].strip()
                    if sl and sl != "[Blank line]":
                        sig_lines.append(sl)
                    i += 1
                if sig_lines:
                    self._add_signature_block(sig_lines)
                continue

            # Preamble text
            if stripped.startswith("The following is information"):
                self._add_body_paragraph(stripped)
                i += 1
                continue

            # Default: body paragraph
            self._add_body_paragraph(stripped)
            i += 1

    # ------------------------------------------------------------------
    # Paragraph helpers
    # ------------------------------------------------------------------

    def _add_formatted_para(self, text: str, fmt: FormatSpec) -> None:
        """Add a paragraph using a FormatSpec from the template."""
        p = self._doc.add_paragraph()
        p.alignment = _ALIGN_MAP.get(fmt.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_before = Twips(fmt.space_before_twips)
        p.paragraph_format.space_after = Twips(fmt.space_after_twips)
        if fmt.line_spacing_twips:
            p.paragraph_format.line_spacing = Twips(fmt.line_spacing_twips)
        if fmt.left_indent_twips:
            p.paragraph_format.left_indent = Twips(fmt.left_indent_twips)
        if fmt.first_line_indent_twips:
            p.paragraph_format.first_line_indent = Twips(fmt.first_line_indent_twips)

        run = p.add_run(text)
        run.font.name = fmt.font_name or self._font
        if fmt.font_size_pt:
            run.font.size = Pt(fmt.font_size_pt)
        run.bold = fmt.bold
        run.italic = fmt.italic

    def _add_body_paragraph(self, text: str) -> None:
        """Regular body text using template body format."""
        fmt = self._t.body_format
        p = self._doc.add_paragraph()
        if fmt.line_spacing_twips:
            p.paragraph_format.line_spacing = Twips(fmt.line_spacing_twips)
        run = p.add_run(text)
        run.font.name = fmt.font_name or self._font
        run.font.size = Pt(fmt.font_size_pt or self._t.default_size_pt)

    def _add_bullet_item(self, text: str) -> None:
        """Summary bullet: indented, bold label if present."""
        fmt = self._t.body_format
        p = self._doc.add_paragraph()
        p.paragraph_format.left_indent = Twips(1080)
        p.paragraph_format.first_line_indent = Twips(-360)
        if fmt.line_spacing_twips:
            p.paragraph_format.line_spacing = Twips(fmt.line_spacing_twips)
        body_size = Pt(fmt.font_size_pt or self._t.default_size_pt)
        if ": " in text:
            label, _, content = text.partition(": ")
            run_label = p.add_run(f"* {label}: ")
            run_label.font.name = self._font
            run_label.font.size = body_size
            run_label.bold = True
            run_content = p.add_run(content)
            run_content.font.name = self._font
            run_content.font.size = body_size
        else:
            run = p.add_run(f"* {text}")
            run.font.name = self._font
            run.font.size = body_size

    def _add_labeled_line(self, label: str, content: str) -> None:
        """Assessment sub-item: bold label + normal content."""
        fmt = self._t.body_format
        body_size = Pt(fmt.font_size_pt or self._t.default_size_pt)
        p = self._doc.add_paragraph()
        if fmt.line_spacing_twips:
            p.paragraph_format.line_spacing = Twips(fmt.line_spacing_twips)
        run_label = p.add_run(label + " ")
        run_label.font.name = self._font
        run_label.font.size = body_size
        run_label.bold = True
        run_content = p.add_run(content)
        run_content.font.name = self._font
        run_content.font.size = body_size

    def _add_memo_line(self, label: str, value: str) -> None:
        """Memo field with tab-aligned value."""
        p = self._doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        body_ls = self._t.body_format.line_spacing_twips
        if body_ls:
            p.paragraph_format.line_spacing = Twips(body_ls)
        pPr = p._element.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "left")
        tab.set(qn("w:pos"), "2160")
        tabs.append(tab)
        pPr.append(tabs)

        run = p.add_run(f"{label}\t{value}")
        run.font.name = self._font
        run.font.size = Pt(self._t.default_size_pt)

    def _add_signature_block(self, lines: list[str]) -> None:
        """Signature block: several empty lines then name, rank, title."""
        for _ in range(5):
            self._doc.add_paragraph()
        for line in lines:
            p = self._doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = self._font
            run.font.size = Pt(self._t.default_size_pt)

    @staticmethod
    def _is_signature_line(text: str) -> bool:
        """Detect signature block lines that come after the conclusion."""
        if text == "[Blank line]":
            return True
        if text.isupper() and len(text.split()) >= 2:
            return True
        if ", " in text and any(
            text.startswith(r) for r in (
                "CPT", "1LT", "2LT", "LT", "MAJ", "COL", "SGM", "MSG",
                "SFC", "SSG", "SGT", "SPC", "PFC", "PV1", "CW",
                "Rank", "OIC", "Commander",
            )
        ):
            return True
        if text.startswith("OIC"):
            return True
        return False
