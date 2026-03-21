"""Converts a TF405 AAR markdown report to a formatted .docx document.

Formatting is derived from the template at:
  AAR_Template/TF405 After Action Report Operation Viper Strike.docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, Twips, Emu


# Inline labels that get bold formatting with normal-weight content after
_INLINE_LABELS = (
    "PERSONNEL:", "EQUIPMENT:", "LOGISTICS:",
    "COMPOSITION:", "STRENGTH:", "CASUALTIES:",
    "Accuracy of Pre-Mission Intelligence:",
    "New Intelligence Gathered During The Operation:",
    "SUSTAINMENT:", "IMPROVEMENT:", "CONCLUSION:",
)

_FONT = "Times New Roman"

# Header banner image path (extracted from template docx)
_BANNER_PATH = Path(__file__).parent / "AAR_Template" / "header_banner.png"

# Regex for section headers: "1. GENERAL INFORMATION ..."
_SECTION_RE = re.compile(r"^(\d+)\.\s+(.+)")

# Regex for memo fields: "TO:    value"
_MEMO_RE = re.compile(r"^(TO|FROM|SUBJECT|REF):\s*(.*)", re.IGNORECASE)


class DocxConverter:
    """Converts a TF405 AAR markdown report to a formatted .docx
    matching the TF405 After Action Report template."""

    def __init__(self, report_text: str) -> None:
        self._lines = report_text.strip().splitlines()
        self._doc = Document()
        self._setup_defaults()

    def save(self, path: str | Path) -> None:
        """Parse the report and save to a .docx file."""
        self._add_header()
        self._build()
        self._add_footer()
        self._doc.save(str(path))

    # ------------------------------------------------------------------
    # Document setup
    # ------------------------------------------------------------------

    def _setup_defaults(self) -> None:
        """Set document-wide defaults matching the template's styles.xml."""
        style = self._doc.styles["Normal"]
        style.font.name = _FONT
        style.font.size = Pt(12)
        # Template docDefaults: sp_before=40 (2pt), sp_after=280 (14pt)
        style.paragraph_format.space_before = Pt(2)
        style.paragraph_format.space_after = Pt(14)
        # Template default line spacing: 259 twips (single)
        style.paragraph_format.line_spacing = Twips(259)

        # Page margins: 1" all sides, header/footer 0.5"
        for section in self._doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)

    # ------------------------------------------------------------------
    # Header
    # ------------------------------------------------------------------

    def _add_header(self) -> None:
        """Add the TF405 banner image to the document header."""
        if not _BANNER_PATH.exists():
            return
        section = self._doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        # Template image: 5731200 x 736600 EMU
        run.add_picture(str(_BANNER_PATH), width=Emu(5731200), height=Emu(736600))

    # ------------------------------------------------------------------
    # Footer
    # ------------------------------------------------------------------

    def _add_footer(self) -> None:
        """Add 'Page X of Y' footer matching the template."""
        section = self._doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = p.add_run("Page ")
        run.font.name = _FONT
        run.font.size = Pt(10)
        run.italic = True

        self._add_field(p, "PAGE", italic=True)

        run = p.add_run(" of ")
        run.font.name = _FONT
        run.font.size = Pt(10)
        run.italic = True

        self._add_field(p, "NUMPAGES", italic=True)

    def _add_field(self, paragraph, field_code: str, italic: bool = False) -> None:
        """Insert a Word field code (PAGE, NUMPAGES, etc.)."""
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run1 = paragraph.add_run()
        run1.font.name = _FONT
        run1.font.size = Pt(10)
        run1.italic = italic
        run1._element.append(fld_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {field_code} "
        run2 = paragraph.add_run()
        run2.font.name = _FONT
        run2.font.size = Pt(10)
        run2.italic = italic
        run2._element.append(instr)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run3 = paragraph.add_run()
        run3.font.name = _FONT
        run3.font.size = Pt(10)
        run3.italic = italic
        run3._element.append(fld_end)

    # ------------------------------------------------------------------
    # Main builder
    # ------------------------------------------------------------------

    def _build(self) -> None:
        """Parse the report lines and build the document."""
        i = 0
        n = len(self._lines)
        in_conclusion = False
        in_header = True   # before section 1
        in_memo = False    # between MEMORANDUM FOR and REF:

        while i < n:
            line = self._lines[i]
            stripped = line.strip()

            # Preserve empty lines as spacing paragraphs (template uses these)
            if not stripped:
                if not in_header or in_memo:
                    self._doc.add_paragraph()
                i += 1
                continue

            # --- Title block ---
            if stripped == "TASK FORCE 405" and i < 3:
                self._add_title_line(stripped, size=Pt(18))
                i += 1
                continue

            if "SWTG" in stripped and "SQUADRON" in stripped:
                self._add_title_line(stripped, size=Pt(14))
                i += 1
                continue

            if ("AFTER ACTION REPORT" in stripped.upper()
                    and not _MEMO_RE.match(stripped)):
                self._add_title_line(stripped, size=Pt(12))
                i += 1
                continue

            # Fallback: "OPERATION" line that got merged with a section header
            # e.g. "OPERATION 1. GENERAL INFORMATION / INTRODUCTION"
            if stripped.startswith("OPERATION "):
                # Use non-anchored search to find "N. SECTION NAME" within the line
                section_in_line = re.search(r"(\d+)\.\s+([A-Z].*)", stripped)
                if section_in_line:
                    # Split: emit the OPERATION part as a title, then the section header
                    op_part = stripped[:section_in_line.start()].strip()
                    sec_part = stripped[section_in_line.start():]
                    if op_part:
                        self._add_title_line(op_part + " - AFTER ACTION REPORT", size=Pt(12))
                    self._add_section_header(sec_part)
                    in_header = False
                else:
                    # Pure OPERATION line without embedded section
                    self._add_title_line(stripped, size=Pt(12))
                i += 1
                continue

            # Date line: "[2025/08/17]" or "2026-03-01" — check BEFORE location
            if re.match(r"^\[?\d{4}[/-]\d{2}[/-]\d{2}\]?$", stripped):
                p = self._doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Twips(480)
                p.paragraph_format.line_spacing = Twips(259)
                run = p.add_run(stripped)
                run.font.name = _FONT
                run.font.size = Pt(12)
                i += 1
                continue

            # Location line (short, uppercase, not a known keyword or section number)
            if (in_header and not in_memo
                    and len(stripped) < 60
                    and not any(c.islower() for c in stripped.replace(" ", ""))
                    and not stripped[0].isdigit()
                    and not stripped.startswith("TASK")
                    and not stripped.startswith("OPERATION")
                    and not stripped.startswith("MEMORANDUM")):
                self._add_title_line(stripped, size=Pt(12))
                i += 1
                continue

            # MEMORANDUM FOR
            if stripped == "MEMORANDUM FOR":
                p = self._doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Twips(400)
                p.paragraph_format.space_after = Twips(120)
                p.paragraph_format.line_spacing = Twips(240)
                run = p.add_run(stripped)
                run.font.name = _FONT
                run.font.size = Pt(18)
                run.bold = True
                self._doc.add_paragraph()  # empty line after
                in_memo = True
                i += 1
                continue

            # Memo fields: TO:, FROM:, SUBJECT:, REF:
            memo_match = _MEMO_RE.match(stripped)
            if memo_match and in_memo:
                label = memo_match.group(1).upper() + ":"
                value = memo_match.group(2).strip()
                self._add_memo_line(label, value)
                if label == "REF:":
                    in_memo = False
                    self._doc.add_paragraph()  # spacing after memo block
                i += 1
                continue

            in_header = False  # past the header/memo block

            # Section headers: "1. GENERAL INFORMATION / INTRODUCTION"
            section_match = _SECTION_RE.match(stripped)
            if section_match:
                section_num = section_match.group(1)
                in_conclusion = section_num == "8"
                self._add_section_header(stripped)
                i += 1
                continue

            # Bullet items: "* Deployed Location: ..."
            if stripped.startswith("* "):
                self._add_bullet_item(stripped[2:])
                i += 1
                continue

            # Inline labels: "PERSONNEL: 0 KIA, 0 WIA, 0 MIA"
            label_found = False
            for label in _INLINE_LABELS:
                if stripped.startswith(label):
                    content = stripped[len(label):].strip()
                    self._add_labeled_line(label, content)
                    label_found = True
                    break
            if label_found:
                i += 1
                continue

            # Signature block detection: lines after conclusion content
            if in_conclusion and self._is_signature_line(stripped):
                sig_lines = []
                while i < n:
                    sl = self._lines[i].strip()
                    if sl and sl != "[Blank line]":
                        sig_lines.append(sl)
                    i += 1
                if sig_lines:
                    self._add_signature_block(sig_lines)
                continue

            # Preamble text (e.g., "The following is information...")
            if stripped.startswith("The following is information"):
                self._add_body_paragraph(stripped)
                i += 1
                continue

            # Default: body paragraph
            self._add_body_paragraph(stripped)
            i += 1

    # ------------------------------------------------------------------
    # Paragraph helpers
    # ------------------------------------------------------------------

    def _add_title_line(self, text: str, size: Pt = Pt(12)) -> None:
        """Title block line: bold, centered, tight line spacing."""
        p = self._doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Twips(276)
        run = p.add_run(text)
        run.font.name = _FONT
        run.font.size = size
        run.bold = True

    def _add_section_header(self, text: str) -> None:
        """Numbered section header with hanging indent."""
        p = self._doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Twips(360)       # 18pt
        p.paragraph_format.line_spacing = Twips(240)
        p.paragraph_format.left_indent = Inches(0.5)      # 720 twips
        p.paragraph_format.first_line_indent = Inches(-0.25)  # hanging 360 twips
        run = p.add_run(text)
        run.font.name = _FONT
        run.font.size = Pt(12)
        run.bold = True

    def _add_body_paragraph(self, text: str) -> None:
        """Regular body text: 11pt, single line spacing."""
        p = self._doc.add_paragraph()
        p.paragraph_format.line_spacing = Twips(259)
        run = p.add_run(text)
        run.font.name = _FONT
        run.font.size = Pt(11)

    def _add_bullet_item(self, text: str) -> None:
        """Summary bullet: indented, bold label if present."""
        p = self._doc.add_paragraph()
        p.paragraph_format.left_indent = Twips(1080)       # 0.75"
        p.paragraph_format.first_line_indent = Twips(-360)  # hanging
        p.paragraph_format.line_spacing = Twips(259)
        if ": " in text:
            label, _, content = text.partition(": ")
            run_label = p.add_run(f"* {label}: ")
            run_label.font.name = _FONT
            run_label.font.size = Pt(11)
            run_label.bold = True
            run_content = p.add_run(content)
            run_content.font.name = _FONT
            run_content.font.size = Pt(11)
        else:
            run = p.add_run(f"* {text}")
            run.font.name = _FONT
            run.font.size = Pt(11)

    def _add_labeled_line(self, label: str, content: str) -> None:
        """Assessment sub-item: bold label + normal content."""
        p = self._doc.add_paragraph()
        p.paragraph_format.line_spacing = Twips(259)
        run_label = p.add_run(label + " ")
        run_label.font.name = _FONT
        run_label.font.size = Pt(11)
        run_label.bold = True
        run_content = p.add_run(content)
        run_content.font.name = _FONT
        run_content.font.size = Pt(11)

    def _add_memo_line(self, label: str, value: str) -> None:
        """Memo field with tab-aligned value."""
        p = self._doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Twips(259)
        # Add tab stop at 1.5" (2160 twips)
        pPr = p._element.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "left")
        tab.set(qn("w:pos"), "2160")
        tabs.append(tab)
        pPr.append(tabs)

        run = p.add_run(f"{label}\t{value}")
        run.font.name = _FONT
        run.font.size = Pt(12)

    def _add_signature_block(self, lines: list[str]) -> None:
        """Signature block: several empty lines then name, rank, title."""
        # Template has ~10 empty paragraphs before signature
        for _ in range(5):
            self._doc.add_paragraph()
        for line in lines:
            p = self._doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = _FONT
            run.font.size = Pt(12)

    @staticmethod
    def _is_signature_line(text: str) -> bool:
        """Detect signature block lines that come after the conclusion."""
        if text == "[Blank line]":
            return True
        if text.isupper() and len(text.split()) >= 2:
            return True
        if ", " in text and any(
            text.startswith(r) for r in (
                "CPT", "1LT", "2LT", "LT", "MAJ", "COL", "SGM", "MSG",
                "SFC", "SSG", "SGT", "SPC", "PFC", "PV1", "CW",
                "Rank", "OIC", "Commander",
            )
        ):
            return True
        if text.startswith("OIC"):
            return True
        return False
